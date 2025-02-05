from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
import json

# Standard section headers for ATS compatibility
STANDARD_SECTIONS = {
    'summary': 'Professional Summary',
    'experience': 'Professional Experience',
    'education': 'Education',
    'skills': 'Techincal Skills',
    'courses': 'Certifications & Training',
    'projects': 'Relevant Projects',
    'hobbies': 'Additional Information'
}

def optimize_content_for_ats(content):
    """Optimize content for ATS parsing by removing special characters and formatting."""
    if isinstance(content, str):
        content = content.replace('\u2022', '')  # Remove bullet points
        content = content.replace('\n', ' ')  # Replace newlines with spaces
        return ' '.join(content.split())  # Remove extra spaces
    return content

def add_heading(document, text, level=1):
    """Add a heading with ATS-friendly formatting."""
    heading = document.add_heading(level=level)
    run = heading.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    paragraph_format = heading.paragraph_format
    paragraph_format.space_before = Pt(4)
    paragraph_format.space_after = Pt(4)
    return heading

def add_name(document, name):
    """Add name section with standard formatting."""
    name_paragraph = document.add_paragraph()
    name_run = name_paragraph.add_run(name.upper())
    name_run.bold = True
    name_run.font.size = Pt(14)
    name_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    reduce_spacing(name_paragraph)

def add_personal_info(document, personal_info):
    """Add contact information section with ATS-friendly formatting."""
    add_heading(document, 'CONTACT INFORMATION')
    for key, value in personal_info.items():
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"{key}: {value}")
        if key.lower() == 'linkedin':
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.underline = True
        reduce_spacing(paragraph)

def add_section(document, title, dates, content, bullet=False, font_size=10):
    """Add a section with standardized formatting."""
    document.add_heading(title, level=2)
    if dates:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(dates)
        run.italic = True
        reduce_spacing(paragraph)

    content = optimize_content_for_ats(content)
    if bullet and isinstance(content, str):
        for point in content.split('. '):
            if point.strip():
                p = document.add_paragraph(style='List Bullet')
                p.add_run(point.strip())
                p.style.font.size = Pt(font_size)
                reduce_spacing(p)
    else:
        paragraph = document.add_paragraph(content)
        paragraph.style.font.size = Pt(font_size)
        reduce_spacing(paragraph)

def add_summary(document, summary):
    """Add professional summary section."""
    add_heading(document, STANDARD_SECTIONS['summary'])
    summary_paragraph = document.add_paragraph(optimize_content_for_ats(summary))
    summary_paragraph.style.font.size = Pt(10)
    reduce_spacing(summary_paragraph)

def add_experience(document, experiences):
    """Add professional experience section with ATS-optimized formatting."""
    add_heading(document, STANDARD_SECTIONS['experience'])
    for exp in experiences:
        exp_title = f"{exp['title']} - {exp['company']}"
        if 'location' in exp:
            exp_title += f" | {exp['location']}"
        exp_dates = exp['dates']
        exp_details = optimize_content_for_ats(exp['description'])
        add_section(document, exp_title, exp_dates, exp_details, bullet=True)

def add_education(document, education):
    """Add education section with standard formatting."""
    add_heading(document, STANDARD_SECTIONS['education'])
    for edu in education:
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{edu['degree']} - {edu['school']}").bold = True
        paragraph.add_run(f" ({edu['dates']})").italic = True
        reduce_spacing(paragraph)

def add_skills(document, skills, section_title='TECHNICAL SKILLS'):
    """Add skills section with categorized formatting."""
    add_heading(document, section_title)
    for skill in skills:
        paragraph = document.add_paragraph(style='List Bullet')
        paragraph.add_run(optimize_content_for_ats(skill))
        paragraph.style.font.size = Pt(10)
        reduce_spacing(paragraph)

def add_courses(document, courses):
    """Add professional development section."""
    add_heading(document, STANDARD_SECTIONS['courses'])
    for course in courses:
        paragraph = document.add_paragraph(style='List Bullet')
        paragraph.add_run(optimize_content_for_ats(course))
        paragraph.style.font.size = Pt(10)
        reduce_spacing(paragraph)

def add_projects(document, projects):
    """Add projects section."""
    add_heading(document, STANDARD_SECTIONS['projects'])
    for project in projects:
        paragraph = document.add_paragraph(style='List Bullet')
        paragraph.add_run(optimize_content_for_ats(project))
        paragraph.style.font.size = Pt(10)
        reduce_spacing(paragraph)

def reduce_spacing(paragraph, line_spacing=Pt(10)):
    """Apply standard spacing to paragraphs."""
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph_format.line_spacing = line_spacing
    paragraph_format.space_after = Pt(2)
    paragraph_format.space_before = Pt(2)

def read_cv_data_from_json(file_path):
    """Read CV data from JSON file."""
    with open(file_path, 'r', encoding='utf-8') as file:
        return json.load(file)

def generate_cv(data):
    """Generate ATS-compliant CV document."""
    document = Document()
    
    # Set standard document properties
    style = document.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    # Set narrow margins
    sections = document.sections
    for section in sections:
        section.top_margin = Pt(36)  # 0.5 inch
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(36)
        section.right_margin = Pt(36)

    # Add document sections
    add_name(document, data['name'])
    add_personal_info(document, data['personal_info'])
    add_summary(document, data['summary'])
    add_experience(document, data['experience'])
    add_education(document, data['education'])
    add_skills(document, data['technical_skills'])
    add_courses(document, data['courses'])
    add_projects(document, data['projects'])
    
    if 'hobbies' in data:
        add_skills(document, data['hobbies'], STANDARD_SECTIONS['hobbies'])

    # Save document with standardized filename
    filename = f"{data['name'].replace(' ', '_')}_Resume.docx"
    document.save(filename)
    return filename

if __name__ == "__main__":
    cv_data = read_cv_data_from_json("cv_data.json")
    output_file = generate_cv(cv_data)
    print(f"CV generated successfully: {output_file}")